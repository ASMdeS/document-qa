# Import necessary classes from the python-docx library
import docx
from docx import Document
from docx.shared import Inches, Pt # For setting margins, font sizes, etc.
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # For text alignment
from docx.text.parfmt import ParagraphFormat # To control spacing
from docx.oxml.ns import qn # For hyperlink style

# --- Helper function to add hyperlinks ---
# (Source: https://github.com/python-openxml/python-docx/issues/74#issuecomment-261169411 - adapted)
def add_hyperlink(paragraph, text, url):
    """
    Adds a hyperlink to a paragraph.

    Args:
        paragraph: The paragraph object to add the hyperlink to.
        text: The display text for the hyperlink.
        url: The URL the hyperlink should point to.

    Returns:
        The hyperlink run object.
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (Word uses 'Hyperlink' style by default)
    r.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1) # Standard blue hyperlink color
    r.font.underline = True

    return r

# --- Content from the provided resume ---
name = "Arthur Santos Marinho de Souza"
email = "arthursmds@gmail.com"
phone = "(81) 98605-6957"
linkedin_url = "https://www.linkedin.com/in/asmdes/"
linkedin_text = linkedin_url # Display text for link

awards_certs_title = "Prêmios e Certificações"
awards_certs_details = [
    "Medalhas: Ouro - Olimpíada Brasileira de Astronomia (2019/2022), Prata - Olimpíada Nacional de Ciências (2022), Prata - Olimpíada Pernambucana de Astronomia (2022), Menção Honrosa - Olimpíada Pernambucana de Física (2022)",
    "Inglês: ECPE (Pass), Duolingo English Test (140/160), SAT (1470/1600)"
]

experience_title = "Experiência"
experience_details = [
    {
        "title": "Planejamento de Supply Chain, Praso – Recife, Pernambuco",
        "duration": "Agosto 2024 – Presente",
        "points": [
            "Reconstruí as ferramentas de previsão de demanda, melhorando a acurácia do nosso planejamento de compras e o risco de estocagem em excesso",
            "Reduzi as perdas por shelf (validade) com um sistema de prevenção de perdas que utiliza prazos de validade e estocagem em excesso, combinados com o sistema de previsão de demanda",
            "Construi uma ferramenta de precificação de produtos, garantindo uma rentabilidade adequada e alinhada com os objetivos da empresa"
        ]
    }
]

organization_title = "Organização"
organization_details = [
     {
        "title": "Membro, UFPE Finance (Liga Acadêmica) – Recife, Pernambuco",
        "duration": "Abril 2024 – Presente",
        "points": [
            "Trainee: Desenvolvi scripts python para realização do valuation intrínseco da Arezzo, analisando fluxo de caixa e balanço patrimonial de forma automatizada, tendo como base teórica os princípios de Aswath Damodaran, além de avaliar o modelo de negócios da empresa.",
            "Itaú Quant: Desenvolvi e validei uma estratégia sistemática de investimentos baseada no indicador KST (Known-Sure Thing), focada na ação TSMC, utilizando dados históricos via API para capturar movimentos de curto prazo e superar o benchmark S&P 500",
            "Constellation Challenge: construí uma análise dos setores de E-Commerce e Fintech na América Latina, para realização de um benchmark para o Mercado Livre"
        ]
    }
]

volunteer_title = "Trabalho Voluntário"
volunteer_details = [
    {
        "title": "Professor de Matemática, Voluntariado Estudantil Marista – Recife, Pernambuco",
        "duration": "Março 2022 – Presente",
        "points": [
            "Ensinei Matemática básica, como as quatro operações, para crianças carentes em idade escolar",
            "Desenvolvi de atividades lúdicas para aprendizado de habilidades socioemocionais",
            "Interagi com as crianças durante os intervalos"
        ]
    }
]

projects_title = "Projetos"
projects_details = [
    {
        "title": "Stock selection performance - Dashboard",
        "link": "https://www.upwork.com/freelancers/~01dd196048866f8814",
        "link_text": "Project Link", # Display text for link
        "points": [
            "Desenvolvi um dashboard dinâmico integrado a planilhas online e sites financeiros, permitindo monitoramento em tempo real de KPIs de desempenho de ações, com feedback máximo do cliente.",
            "Implementei um módulo de backtesting utilizando dados de 2023, identificando padrões de desempenho e otimizando estratégias de seleção de ações.",
            "Criei uma interface intuitiva com filtros personalizados e documentei todo o sistema, garantindo usabilidade e suporte pós-implementação eficiente."
        ]
    }
]

education_title = "Formação Acadêmica"
education_details = [
    "Universidade Federal de Pernambuco (UFPE) – Bacharelado em Sistemas de Informação\t  2023.1 - 2026.2"
]

# --- Create the DOCX document ---
doc = Document()

# --- Set Margins (Optional, can help fit content) ---
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# --- Define paragraph spacing ---
# Function to apply spacing settings to a paragraph
def set_paragraph_spacing(paragraph, before=Pt(6), after=Pt(2), line_spacing=1.0):
     fmt = paragraph.paragraph_format
     fmt.space_before = before
     fmt.space_after = after
     # fmt.line_spacing = line_spacing # Uncomment to adjust line spacing if needed

# --- Add Name and Contact Info ---
name_paragraph = doc.add_paragraph()
name_run = name_paragraph.add_run(name)
name_run.bold = True
name_run.font.size = Pt(14) # Slightly smaller font size
name_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
set_paragraph_spacing(name_paragraph, before=Pt(0), after=Pt(2)) # Reduced spacing

# Add contact info with hyperlink
contact_paragraph = doc.add_paragraph()
contact_paragraph.add_run(f"{email} | {phone} | ")
add_hyperlink(contact_paragraph, linkedin_text, linkedin_url)
contact_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
set_paragraph_spacing(contact_paragraph, before=Pt(0), after=Pt(6)) # Spacing after contact


# --- Add Awards and Certifications ---
heading_awards = doc.add_heading(awards_certs_title, level=1)
set_paragraph_spacing(heading_awards, before=Pt(8), after=Pt(2)) # Spacing for heading
for item in awards_certs_details:
    p = doc.add_paragraph(item, style='List Bullet')
    set_paragraph_spacing(p, before=Pt(0), after=Pt(2)) # Spacing for list items
# Removed extra spacing paragraph

# --- Add Experience ---
heading_exp = doc.add_heading(experience_title, level=1)
set_paragraph_spacing(heading_exp, before=Pt(8), after=Pt(2))
for job in experience_details:
    p_title = doc.add_paragraph()
    p_title.add_run(job['title']).bold = True
    p_title.add_run(f"\t {job['duration']}")
    set_paragraph_spacing(p_title, before=Pt(4), after=Pt(0)) # Spacing for job title

    for point in job['points']:
        p = doc.add_paragraph(point, style='List Bullet')
        set_paragraph_spacing(p, before=Pt(0), after=Pt(2))
# Removed extra spacing paragraph

# --- Add Organization ---
heading_org = doc.add_heading(organization_title, level=1)
set_paragraph_spacing(heading_org, before=Pt(8), after=Pt(2))
for org in organization_details:
    p_title = doc.add_paragraph()
    p_title.add_run(org['title']).bold = True
    p_title.add_run(f"\t {org['duration']}")
    set_paragraph_spacing(p_title, before=Pt(4), after=Pt(0))

    for point in org['points']:
        if ":" in point:
             parts = point.split(":", 1)
             p_bullet = doc.add_paragraph(style='List Bullet')
             p_bullet.add_run(parts[0] + ":").bold = True
             p_bullet.add_run(parts[1].strip()) # Use strip() to remove leading space
             set_paragraph_spacing(p_bullet, before=Pt(0), after=Pt(2))
        else:
            p = doc.add_paragraph(point, style='List Bullet')
            set_paragraph_spacing(p, before=Pt(0), after=Pt(2))
# Removed extra spacing paragraph

# --- Add Volunteer Work ---
heading_vol = doc.add_heading(volunteer_title, level=1)
set_paragraph_spacing(heading_vol, before=Pt(8), after=Pt(2))
for volunteer in volunteer_details:
    p_title = doc.add_paragraph()
    p_title.add_run(volunteer['title']).bold = True
    p_title.add_run(f"\t {volunteer['duration']}")
    set_paragraph_spacing(p_title, before=Pt(4), after=Pt(0))

    for point in volunteer['points']:
        p = doc.add_paragraph(point, style='List Bullet')
        set_paragraph_spacing(p, before=Pt(0), after=Pt(2))
# Removed extra spacing paragraph

# --- Add Projects ---
heading_proj = doc.add_heading(projects_title, level=1)
set_paragraph_spacing(heading_proj, before=Pt(8), after=Pt(2))
for project in projects_details:
    p_title = doc.add_paragraph()
    p_title.add_run(project['title']).bold = True
    if 'link' in project and 'link_text' in project: # Add hyperlink if present
         p_title.add_run("\t") # Add a tab for spacing
         add_hyperlink(p_title, project['link_text'], project['link'])
    set_paragraph_spacing(p_title, before=Pt(4), after=Pt(0))

    for point in project['points']:
        p = doc.add_paragraph(point, style='List Bullet')
        set_paragraph_spacing(p, before=Pt(0), after=Pt(2))
# Removed extra spacing paragraph

# --- Add Education ---
heading_edu = doc.add_heading(education_title, level=1)
set_paragraph_spacing(heading_edu, before=Pt(8), after=Pt(2))
for edu in education_details:
    p = doc.add_paragraph(edu)
    set_paragraph_spacing(p, before=Pt(0), after=Pt(2))

# --- Save the document ---
file_name = 'Arthur_Souza_Resume_Updated.docx'
try:
    doc.save(file_name)
    print(f"Document '{file_name}' created successfully.")
except Exception as e:
    print(f"Error saving document: {e}")

